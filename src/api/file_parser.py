# Copyright (c) Microsoft. All rights reserved.
# Licensed under the MIT license. See LICENSE file in the project root for full license information.

import os
import json
import logging
from typing import Dict, List, Optional, Any, Tuple
import io

logger = logging.getLogger("azureaiapp")


class FileParser:
    """
    Parser for different file formats to extract content and metadata.
    This class provides methods to parse various file formats for use in search indexing.
    """

    @staticmethod
    def parse_file(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a file based on its extension and return content and metadata.

        :param file_path: Path to the file to parse
        :return: Tuple of (content, metadata)
        """
        extension = os.path.splitext(file_path)[1].lower()

        if extension == '.json':
            return FileParser._parse_json(file_path)
        elif extension == '.md':
            return FileParser._parse_markdown(file_path)
        elif extension == '.txt':
            return FileParser._parse_text(file_path)
        elif extension == '.pdf':
            return FileParser._parse_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return FileParser._parse_docx(file_path)
        else:
            logger.warning(f"Unsupported file format: {extension}")
            return "", {"error": f"Unsupported file format: {extension}"}

    @staticmethod
    def _parse_json(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse JSON file and extract content and metadata"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)

            # Extract metadata
            metadata = {
                "file_type": "json",
                "file_name": os.path.basename(file_path),
                "categories": []
            }

            # Look for specific metadata fields in JSON structure
            if isinstance(data, dict):
                if "category" in data:
                    metadata["categories"].append(data["category"])
                if "brand" in data:
                    metadata["brand"] = data["brand"]
                if "id" in data:
                    metadata["document_id"] = data["id"]

            # Convert JSON to text for indexing
            content = json.dumps(data, indent=2)

            return content, metadata
        except Exception as e:
            logger.error(f"Error parsing JSON file {file_path}: {str(e)}")
            return "", {"error": str(e), "file_name": os.path.basename(file_path)}

    @staticmethod
    def _parse_markdown(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse Markdown file and extract content and metadata"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            # Extract metadata
            metadata = {
                "file_type": "markdown",
                "file_name": os.path.basename(file_path),
                "categories": []
            }

            # Look for headers and extract metadata from them
            lines = content.split('\n')
            for i, line in enumerate(lines):
                if line.startswith('# '):
                    metadata["title"] = line[2:].strip()
                elif line.startswith('## Category'):
                    if i+1 < len(lines) and not lines[i+1].startswith('#'):
                        category = lines[i+1].strip()
                        metadata["categories"].append(category)
                elif line.startswith('## Brand'):
                    if i+1 < len(lines) and not lines[i+1].startswith('#'):
                        metadata["brand"] = lines[i+1].strip()

            # Look for product ID in the filename or content
            if "item_number:" in content:
                parts = content.split("item_number:")
                if len(parts) > 1:
                    item_number = parts[1].split("\n")[0].strip()
                    metadata["document_id"] = item_number

            return content, metadata
        except Exception as e:
            logger.error(f"Error parsing Markdown file {file_path}: {str(e)}")
            return "", {"error": str(e), "file_name": os.path.basename(file_path)}

    @staticmethod
    def _parse_text(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse plain text file and extract content and metadata"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            # Extract metadata (limited for text files)
            metadata = {
                "file_type": "text",
                "file_name": os.path.basename(file_path),
                "categories": []
            }

            return content, metadata
        except Exception as e:
            logger.error(f"Error parsing text file {file_path}: {str(e)}")
            return "", {"error": str(e), "file_name": os.path.basename(file_path)}

    @staticmethod
    def _parse_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse PDF file and extract content and metadata"""
        try:
            # Import here to avoid dependency issues if not installed
            try:
                import PyPDF2
            except ImportError:
                logger.error(
                    "PyPDF2 library not installed. Please install it to parse PDF files.")
                return "", {"error": "PyPDF2 library not installed", "file_name": os.path.basename(file_path)}

            content = ""
            metadata = {
                "file_type": "pdf",
                "file_name": os.path.basename(file_path),
                "categories": []
            }

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract metadata from PDF info
                if pdf_reader.metadata:
                    for key, value in pdf_reader.metadata.items():
                        if key == '/Title':
                            metadata["title"] = value
                        elif key == '/Subject':
                            metadata["subject"] = value
                        elif key == '/Keywords':
                            keywords = value.split(',')
                            metadata["categories"] = [k.strip()
                                                      for k in keywords]

                # Extract content from all pages
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n\n"

            return content, metadata
        except Exception as e:
            logger.error(f"Error parsing PDF file {file_path}: {str(e)}")
            return "", {"error": str(e), "file_name": os.path.basename(file_path)}

    @staticmethod
    def _parse_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse DOCX file and extract content and metadata"""
        try:
            # Import here to avoid dependency issues if not installed
            try:
                import docx
            except ImportError:
                logger.error(
                    "python-docx library not installed. Please install it to parse DOCX files.")
                return "", {"error": "python-docx library not installed", "file_name": os.path.basename(file_path)}

            doc = docx.Document(file_path)
            content = "\n".join(
                [paragraph.text for paragraph in doc.paragraphs])

            # Extract metadata
            metadata = {
                "file_type": "docx",
                "file_name": os.path.basename(file_path),
                "categories": []
            }

            # Try to extract title from first paragraph if it looks like a title
            if doc.paragraphs and doc.paragraphs[0].text.strip():
                first_para = doc.paragraphs[0].text.strip()
                if len(first_para) < 100:  # Assume short first paragraphs are titles
                    metadata["title"] = first_para

            # Extract core properties
            try:
                metadata["title"] = doc.core_properties.title or metadata.get(
                    "title", "")
                metadata["author"] = doc.core_properties.author
                if doc.core_properties.keywords:
                    keywords = doc.core_properties.keywords.split(',')
                    metadata["categories"] = [k.strip() for k in keywords]
            except:
                pass

            return content, metadata
        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {str(e)}")
            return "", {"error": str(e), "file_name": os.path.basename(file_path)}

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Chunk text into semantically meaningful segments.

        :param text: Text to chunk
        :param chunk_size: Target size for each chunk
        :param overlap: Overlap between chunks to maintain context
        :return: List of text chunks
        """
        if not text:
            return []

        # If text is smaller than chunk size, return as is
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            # Get a chunk of text
            end = start + chunk_size

            if end >= len(text):
                # Last chunk
                chunk = text[start:]
                chunks.append(chunk)
                break

            # Try to find a good breaking point (e.g., paragraph break)
            break_point = text.rfind('\n\n', start, end)

            if break_point == -1:
                # Try to find a sentence break
                break_point = text.rfind('. ', start, end)

            if break_point == -1:
                # No good break point found, just use the chunk size
                break_point = end
            else:
                # Include the break character(s)
                break_point += 2

            chunk = text[start:break_point]
            chunks.append(chunk)

            # Start next chunk with overlap for context
            start = break_point - overlap if break_point > overlap else break_point

        return chunks

    @staticmethod
    def format_citation(metadata: Dict[str, Any], section: str = "") -> str:
        """
        Format citation information based on metadata.

        :param metadata: File metadata
        :param section: Optional section or part of the document being cited
        :return: Formatted citation string
        """
        file_name = metadata.get("file_name", "Unknown source")
        title = metadata.get("title", file_name)

        citation = f"Source: {title}"

        if section:
            citation += f", Section: {section}"

        if "author" in metadata:
            citation += f", Author: {metadata['author']}"

        if "brand" in metadata:
            citation += f", Brand: {metadata['brand']}"

        if "document_id" in metadata:
            citation += f", ID: {metadata['document_id']}"

        return citation
